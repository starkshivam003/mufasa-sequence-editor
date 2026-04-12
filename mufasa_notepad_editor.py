import tkinter as tk
from tkinter import filedialog, colorchooser, messagebox
import os

# Try to import docx, but don't crash if the user hasn't installed it yet
try:
    from docx import Document
    from docx.shared import RGBColor, Pt, Inches
    from docx.enum.section import WD_ORIENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

class SequenceNotepad:
    def __init__(self, root):
        self.root = root
        # self.root.title("Sequence Highlighter")
        # self.root.geometry("1200x800")
        # Text area setup (Courier font ensures alignments are visible while editing)
        self.text = tk.Text(root, wrap="none", undo=True, font=("Courier New", 12))
        self.text.pack(expand=1, fill="both")

        # Scrollbars (crucial for long FASTA/Clustal lines)
        scroll_y = tk.Scrollbar(self.text, orient="vertical", command=self.text.yview)
        scroll_x = tk.Scrollbar(self.text, orient="horizontal", command=self.text.xview)
        self.text.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        scroll_y.pack(side="right", fill="y")
        scroll_x.pack(side="bottom", fill="x")

        # Menu
        # menu = tk.Menu(root)
        # root.config(menu=menu)
        top_level = root.winfo_toplevel() 
        menu = tk.Menu(top_level)
        top_level.config(menu=menu)

        # File menu
        file_menu = tk.Menu(menu, tearoff=0)
        menu.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Open File...", command=self.open_file)
        file_menu.add_separator()
        
        # Save As Submenu
        save_menu = tk.Menu(file_menu, tearoff=0)
        file_menu.add_cascade(label="Save As...", menu=save_menu)
        save_menu.add_command(label="Plain Text (.txt / .fasta)", command=self.save_txt)
        save_menu.add_command(label="Word Document (.docx)", command=self.save_docx)
        save_menu.add_command(label="Rich Text Format (.rtf)", command=self.save_rtf)
        
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=root.quit)

        # Highlight menu
        highlight_menu = tk.Menu(menu, tearoff=0)
        menu.add_cascade(label="Highlight", menu=highlight_menu)
        highlight_menu.add_command(label="Highlight Selection", command=self.highlight_text)
        highlight_menu.add_command(label="Remove Highlight (Selection)", command=self.remove_selected_highlight)

        self.tag_count = 0
        self.root.bind("<Control-f>", self.open_find_window)
        self.last_pos = "1.0"

    # ---------- INPUT ENGINE ----------
    def open_file(self):
        file_path = filedialog.askopenfilename()
        if file_path:
            with open(file_path, "r") as f:
                content = f.read()
            self.text.delete(1.0, tk.END)
            # Insert exactly as it was written in the file
            self.text.insert(tk.END, content)
            self.last_pos = "1.0"

    # ---------- PARSING ENGINE ----------
    def _get_colored_segments(self):
        """
        Scans the text widget and breaks it down into chunks of (string, hex_color).
        This guarantees we don't miss any line breaks or overlap colors.
        """
        dump = self.text.dump("1.0", "end-1c", text=True, tag=True)
        active_tags = []
        segments = []

        for type_, value, index in dump:
            if type_ == "tagon" and value.startswith("highlight"):
                active_tags.append(value)
            elif type_ == "tagoff" and value.startswith("highlight"):
                if value in active_tags:
                    active_tags.remove(value)
            elif type_ == "text":
                color = None
                if active_tags:
                    top_tag = active_tags[-1]
                    color = self.text.tag_cget(top_tag, "foreground")
                segments.append((value, color))
                
        return segments

    # ---------- OUTPUT ENGINES ----------
    def save_txt(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".txt", 
                                                 filetypes=[("Text/FASTA Files", "*.txt *.fasta *.clustal"), ("All Files", "*.*")])
        if not file_path:
            return
            
        content = self.text.get("1.0", "end-1c")
        with open(file_path, "w") as f:
            f.write(content)
        messagebox.showinfo("Success", "Saved safely without colors.")

    def save_docx(self):
        if not HAS_DOCX:
            messagebox.showerror("Missing Library", "Please run 'pip install python-docx' in your terminal to save as Word.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path:
            return

        doc = Document()
        
        # --- NEW ARCHITECTURE: PAGE LAYOUT ---
        section = doc.sections[-1]
        
        # 1. Force Landscape orientation
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        
        # 2. Force Narrow Margins (0.5 inches)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        # 3. Setup Monospace Font and Size
        style = doc.styles['Normal']
        style.font.name = 'Courier New'
        style.font.size = Pt(8)  # Reduced to 8pt to fit long Clustal lines
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

        segments = self._get_colored_segments()

        for text, color in segments:
            parts = text.split('\n')
            for i, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    if color:
                        r, g, b = int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
                
                if i < len(parts) - 1:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)

        doc.save(file_path)
        messagebox.showinfo("Success", "Saved Word Document with Landscape alignment.")

    def save_rtf(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".rtf", filetypes=[("Rich Text Format", "*.rtf")])
        if not file_path:
            return

        segments = self._get_colored_segments()
        unique_colors = set(c for t, c in segments if c)
        
        color_to_idx = {c: i+1 for i, c in enumerate(unique_colors)}

        # RTF Header
        rtf = ["{\\rtf1\\ansi\\deff0"]
        rtf.append("{\\fonttbl{\\f0\\fmodern\\fcharset0 Courier New;}}")

        # Color Table
        if unique_colors:
            ctbl = "{\\colortbl;"
            for c in unique_colors:
                r, g, b = int(c[1:3], 16), int(c[3:5], 16), int(c[5:7], 16)
                ctbl += f"\\red{r}\\green{g}\\blue{b};"
            ctbl += "}"
            rtf.append(ctbl)

        # --- NEW ARCHITECTURE: PAGE LAYOUT FOR RTF ---
        # \landscape = force landscape orientation
        # \paperw15840\paperh12240 = 11x8.5 inches (RTF uses "twips", 1440 twips = 1 inch)
        # \margl720\margr720... = 0.5 inch margins on all sides (720 twips)
        rtf.append("\\landscape\\paperw15840\\paperh12240\\margl720\\margr720\\margt720\\margb720\n")

        # \fs16 = 8pt font (RTF measures fonts in half-points, so 16 = 8pt)
        rtf.append("\\f0\\fs16\n")

        # Write text and color codes
        for text, color in segments:
            # Escape RTF special characters so they don't break the file
            text = text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
            text = text.replace('\n', '\\par\n')

            if color:
                idx = color_to_idx[color]
                rtf.append(f"\\cf{idx} {text}\\cf0 ")
            else:
                rtf.append(text)

        rtf.append("}")

        with open(file_path, "w") as f:
            f.write("".join(rtf))
            
        messagebox.showinfo("Success", "Saved as Rich Text Format with Landscape alignment.")

    # ---------- HIGHLIGHT LOGIC ----------
    def highlight_text(self):
        try:
            start = self.text.index("sel.first")
            end = self.text.index("sel.last")
        except tk.TclError:
            return

        color = colorchooser.askcolor(title="Choose Motif Color")[1]
        if not color:
            return

        tag_name = f"highlight{self.tag_count}"
        self.tag_count += 1

        self.text.tag_add(tag_name, start, end)
        self.text.tag_config(tag_name, foreground=color) # Changed to foreground for better readability

    def remove_selected_highlight(self):
        try:
            start = self.text.index("sel.first")
            end = self.text.index("sel.last")
        except tk.TclError:
            return

        for tag in self.text.tag_names():
            if tag.startswith("highlight"):
                self.text.tag_remove(tag, start, end)

    # ---------- FIND WINDOW ----------
    def open_find_window(self, event=None):
        self.find_window = tk.Toplevel(self.root)
        self.find_window.title("Find Sequence")
        self.find_window.geometry("250x50")

        tk.Label(self.find_window, text="Find:").pack(side="left", padx=5)

        self.find_entry = tk.Entry(self.find_window)
        self.find_entry.pack(side="left", fill="both", expand=1)
        self.find_entry.focus()

        tk.Button(self.find_window, text="Next", command=self.find_next).pack(side="left", padx=5)
        self.find_window.protocol("WM_DELETE_WINDOW", self.close_find_window)

    def close_find_window(self):
        self.text.tag_remove("search", "1.0", tk.END)
        self.find_window.destroy()

    def find_next(self):
        search_term = self.find_entry.get()
        if not search_term:
            return

        self.text.tag_remove("search", "1.0", tk.END)
        pos = self.text.search(search_term, self.last_pos, stopindex=tk.END)

        if not pos:
            self.last_pos = "1.0"
            return

        end = f"{pos}+{len(search_term)}c"
        
        self.text.tag_add("search", pos, end)
        self.text.tag_config("search", background="yellow")

        self.text.tag_remove("sel", "1.0", tk.END)
        self.text.tag_add("sel", pos, end)
        self.text.mark_set(tk.INSERT, end)
        self.text.see(pos)

        self.last_pos = end


if __name__ == "__main__":
    root = tk.Tk()
    app = SequenceNotepad(root)
    root.mainloop()
