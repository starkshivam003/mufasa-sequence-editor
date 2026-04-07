import tkinter as tk
from tkinter import messagebox, filedialog
import re
import webbrowser
import os

# Try to import python-docx for Word exports
try:
    from docx import Document
    from docx.shared import RGBColor, Pt, Inches
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_COLOR_INDEX
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

class MufasaV2:
    def __init__(self, root):
        self.root = root
        self.root.title("MUFASA V2 - Sequence Automation & Heatmap Engine")
        self.root.geometry("1100x680")

        # --- UI LAYOUT ---
        main_frame = tk.Frame(root, padx=10, pady=10)
        main_frame.pack(expand=True, fill="both")

        # Left Panel (Parent Sequence)
        left_frame = tk.Frame(main_frame)
        left_frame.pack(side="left", expand=True, fill="both", padx=(0, 5))
        tk.Label(left_frame, text="1. Paste Parent Sequence (FASTA or raw text):", font=("Arial", 10, "bold")).pack(anchor="w")
        self.seq_text = tk.Text(left_frame, wrap="word", font=("Courier New", 10))
        self.seq_text.pack(expand=True, fill="both")

        # Right Panel (Peptides)
        right_frame = tk.Frame(main_frame)
        right_frame.pack(side="right", expand=True, fill="both", padx=(5, 0))
        tk.Label(right_frame, text="2. Paste Peptide Column (from Excel):", font=("Arial", 10, "bold")).pack(anchor="w")
        self.pep_text = tk.Text(right_frame, wrap="none", font=("Courier New", 10))
        self.pep_text.pack(expand=True, fill="both")

        # Options Panel (Reverse Mode)
        options_frame = tk.Frame(root, padx=10)
        options_frame.pack(fill="x")
        self.reverse_mode = tk.BooleanVar(value=False)
        tk.Checkbutton(options_frame, text="Reverse Mode (Highlight missing gaps, hide mapped regions)", 
                       variable=self.reverse_mode, font=("Arial", 10, "bold"), fg="#D32F2F").pack(side="left", pady=5)

        # Bottom Button
        btn_frame = tk.Frame(root, pady=10)
        btn_frame.pack(fill="x")
        generate_btn = tk.Button(btn_frame, text="GENERATE COVERAGE MAP", font=("Arial", 12, "bold"), bg="#4CAF50", fg="white", command=self.generate_preview)
        generate_btn.pack(pady=10, ipadx=20, ipady=5)

        # Color Theory Rulebook: Jewel Tones for high contrast
        self.font_palette = ["#000080", "#800000", "#006400", "#4B0082", "#8B4513", "#2F4F4F"]

    # ---------- PHASE 1: SANITIZATION LAYER ----------
    def clean_peptide(self, raw_pep):
        """Strips cleavage dots and PTM brackets (e.g., K.A[+80]SDF.R -> ASDF)"""
        pep = raw_pep.strip()
        if not pep: return ""
        match = re.search(r"\.(.*?)\.", pep)
        if match:
            pep = match.group(1)
        pep = re.sub(r'[^a-zA-Z]', '', pep)
        return pep.upper()

    def clean_sequence(self, raw_seq):
        """Removes FASTA headers and all hidden line breaks/spaces."""
        lines = raw_seq.split('\n')
        clean_lines = [line for line in lines if not line.startswith('>')]
        clean_str = "".join(clean_lines)
        clean_str = re.sub(r'\s+', '', clean_str)
        return clean_str.upper()

    # ---------- PHASE 2: DUAL-ARRAY MAPPING ENGINE ----------
    def generate_preview(self):
        raw_seq = self.seq_text.get("1.0", tk.END)
        raw_peps = self.pep_text.get("1.0", tk.END).split('\n')

        sequence = self.clean_sequence(raw_seq)
        if not sequence:
            messagebox.showwarning("Empty Input", "Please provide a parent sequence.")
            return

        peptides = list(set(filter(None, [self.clean_peptide(p) for p in raw_peps])))
        if not peptides:
            messagebox.showwarning("Empty Input", "Please provide at least one valid peptide.")
            return

        peptides.sort(key=len)

        font_colors = [None] * len(sequence)
        hit_counts = [0] * len(sequence)
        unmapped = []

        # Standard Mapping Layer
        for idx, pep in enumerate(peptides):
            color = self.font_palette[idx % len(self.font_palette)]
            start = 0
            found = False
            
            while True:
                start = sequence.find(pep, start)
                if start == -1: break
                found = True
                for i in range(start, start + len(pep)):
                    font_colors[i] = color
                    hit_counts[i] += 1
                start += 1 
            
            if not found:
                unmapped.append(pep)

        # --- THE REVERSE LOGIC FLIP ---
        if self.reverse_mode.get():
            for i in range(len(sequence)):
                if hit_counts[i] > 0:
                    # Mapped region: Strip all colors and heatmap background
                    font_colors[i] = None
                    hit_counts[i] = 0
                else:
                    # Unmapped gap: Apply strong red font to flag it
                    font_colors[i] = "#D32F2F"
                    hit_counts[i] = 1 # Sets it to bold, but prevents heatmap background

        self.open_preview_window(sequence, font_colors, hit_counts, unmapped)

    # ---------- PHASE 3: PREVIEW & CLUSTAL WRAP ----------
    def get_bg_color(self, hits):
        """Calculates Heatmap Gradient based on coverage depth"""
        if hits <= 1: return None
        if hits == 2: return "#FFF9C4" # Soft Yellow
        if hits == 3: return "#FFE0B2" # Soft Orange
        return "#FFCDD2"               # Soft Red (4+ hits)

    def open_preview_window(self, sequence, font_colors, hit_counts, unmapped):
        preview_win = tk.Toplevel(self.root)
        preview_win.title("MUFASA Coverage Map & Export")
        preview_win.geometry("950x700")

        # Orphan Report
        if unmapped and not self.reverse_mode.get():
            warning_text = f"⚠️ WARNING: {len(unmapped)} peptides could not be mapped to the parent sequence."
            tk.Label(preview_win, text=warning_text, fg="#D32F2F", font=("Arial", 10, "bold")).pack(pady=5)

        # Legend (Only show if Reverse Mode is OFF)
        if not self.reverse_mode.get():
            legend_frame = tk.Frame(preview_win)
            legend_frame.pack(pady=5)
            tk.Label(legend_frame, text="Heatmap Depth: ", font=("Arial", 9, "bold")).pack(side="left")
            tk.Label(legend_frame, text=" 1 Hit ", font=("Arial", 9)).pack(side="left", padx=2)
            tk.Label(legend_frame, text=" 2 Hits ", bg="#FFF9C4", font=("Arial", 9)).pack(side="left", padx=2)
            tk.Label(legend_frame, text=" 3 Hits ", bg="#FFE0B2", font=("Arial", 9)).pack(side="left", padx=2)
            tk.Label(legend_frame, text=" 4+ Hits ", bg="#FFCDD2", font=("Arial", 9)).pack(side="left", padx=2)
        else:
            tk.Label(preview_win, text="🔴 REVERSE MODE ACTIVE: Displaying Missing Sequences", fg="#D32F2F", font=("Arial", 10, "bold")).pack(pady=10)

        preview_text = tk.Text(preview_win, wrap="none", font=("Courier New", 12), bg="#FAFAFA")
        preview_text.pack(expand=True, fill="both", padx=10, pady=5)

        # --- CLUSTAL WRAP ENGINE (60 chars per line) ---
        chunk_size = 60
        for i in range(0, len(sequence), chunk_size):
            chunk_seq = sequence[i:i+chunk_size]
            chunk_fg = font_colors[i:i+chunk_size]
            chunk_hits = hit_counts[i:i+chunk_size]

            current_str = chunk_seq[0]
            current_fg = chunk_fg[0]
            current_bg = self.get_bg_color(chunk_hits[0])

            for j in range(1, len(chunk_seq)):
                bg = self.get_bg_color(chunk_hits[j])
                fg = chunk_fg[j]

                if fg == current_fg and bg == current_bg:
                    current_str += chunk_seq[j]
                else:
                    self._insert_styled(preview_text, current_str, current_fg, current_bg)
                    current_str = chunk_seq[j]
                    current_fg = fg
                    current_bg = bg
            
            self._insert_styled(preview_text, current_str, current_fg, current_bg)
            preview_text.insert(tk.END, "\n")

        preview_text.config(state=tk.DISABLED)

        # Action Bar
        btn_frame = tk.Frame(preview_win, pady=10)
        btn_frame.pack()
        tk.Button(btn_frame, text="Export HTML", command=lambda: self.save_html(preview_text)).pack(side="left", padx=5)
        tk.Button(btn_frame, text="Export Word (.docx)", command=lambda: self.save_docx(preview_text)).pack(side="left", padx=5)
        tk.Button(btn_frame, text="Export RTF (.rtf)", command=lambda: self.save_rtf(preview_text)).pack(side="left", padx=5)
        tk.Button(btn_frame, text="Close Preview", command=preview_win.destroy).pack(side="left", padx=5)

    def _insert_styled(self, widget, text, fg, bg):
        fg_str = fg if fg else "NONE"
        bg_str = bg if bg else "NONE"
        tag = f"style|{fg_str}|{bg_str}"
        
        kwargs = {"font": ("Courier New", 12, "bold" if fg else "normal")}
        if fg: kwargs["foreground"] = fg
        if bg: kwargs["background"] = bg
            
        widget.tag_configure(tag, **kwargs)
        widget.insert(tk.END, text, tag)

    # ---------- PHASE 4: EXPORT ENGINES ----------
    def _extract_segments(self, text_widget):
        dump = text_widget.dump("1.0", "end-1c", text=True, tag=True)
        active_fg = None
        active_bg = None
        segments = []

        for type_, value, index in dump:
            if type_ == "tagon" and value.startswith("style|"):
                parts = value.split("|")
                active_fg = parts[1] if parts[1] != "NONE" else None
                active_bg = parts[2] if parts[2] != "NONE" else None
            elif type_ == "tagoff" and value.startswith("style|"):
                active_fg = None
                active_bg = None
            elif type_ == "text":
                segments.append((value, active_fg, active_bg))
        return segments

    def save_html(self, text_widget):
        file_path = filedialog.asksaveasfilename(defaultextension=".html", filetypes=[("HTML File", "*.html")])
        if not file_path: return

        segments = self._extract_segments(text_widget)
        html = ["<html><body style='background:#ffffff; color:#333333; padding:20px;'>"]
        title = "MUFASA Missing Sequences Map" if self.reverse_mode.get() else "MUFASA Depth Heatmap"
        html.append(f"<h2>{title}</h2>")
        html.append("<pre style='font-family:\"Courier New\", Courier, monospace; font-size:14px; line-height:1.8;'>")

        for text, fg, bg in segments:
            style = ""
            if fg: style += f"color:{fg}; font-weight:bold; "
            if bg: style += f"background-color:{bg}; "
            
            if style: html.append(f"<span style='{style}'>{text}</span>")
            else: html.append(text)

        html.append("</pre></body></html>")
        with open(file_path, "w", encoding="utf-8") as f: f.write("".join(html))
        webbrowser.open(f"file://{os.path.abspath(file_path)}")

    def save_docx(self, text_widget):
        if not HAS_DOCX:
            messagebox.showerror("Missing Library", "Please run 'pip install python-docx'.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path: return

        bg_to_highlight = {
            "#FFF9C4": WD_COLOR_INDEX.YELLOW,
            "#FFE0B2": WD_COLOR_INDEX.DARK_YELLOW, 
            "#FFCDD2": WD_COLOR_INDEX.PINK
        }

        doc = Document()
        section = doc.sections[-1]
        new_w, new_h = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = new_w, new_h
        section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0.5)

        style = doc.styles['Normal']
        style.font.name = 'Courier New'
        style.font.size = Pt(9)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

        segments = self._extract_segments(text_widget)
        for text, fg, bg in segments:
            parts = text.split('\n')
            for i, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    if fg:
                        run.font.color.rgb = RGBColor(int(fg[1:3], 16), int(fg[3:5], 16), int(fg[5:7], 16))
                        run.font.bold = True
                    if bg and bg in bg_to_highlight:
                        run.font.highlight_color = bg_to_highlight[bg]
                
                if i < len(parts) - 1:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)

        doc.save(file_path)
        messagebox.showinfo("Success", "Saved Word Document.")

    def save_rtf(self, text_widget):
        file_path = filedialog.asksaveasfilename(defaultextension=".rtf", filetypes=[("Rich Text", "*.rtf")])
        if not file_path: return

        segments = self._extract_segments(text_widget)
        
        unique_fgs = set(fg for t, fg, bg in segments if fg)
        unique_bgs = set(bg for t, fg, bg in segments if bg)
        
        color_tbl_list = list(unique_fgs) + list(unique_bgs)
        color_idx = {c: i+1 for i, c in enumerate(color_tbl_list)}

        rtf = ["{\\rtf1\\ansi\\deff0{\\fonttbl{\\f0\\fmodern\\fcharset0 Courier New;}}"]
        
        if color_tbl_list:
            ctbl = "{\\colortbl;"
            for c in color_tbl_list:
                r, g, b = int(c[1:3], 16), int(c[3:5], 16), int(c[5:7], 16)
                ctbl += f"\\red{r}\\green{g}\\blue{b};"
            ctbl += "}"
            rtf.append(ctbl)

        rtf.append("\\landscape\\paperw15840\\paperh12240\\margl720\\margr720\\margt720\\margb720\n")
        rtf.append("\\f0\\fs18\n")

        for text, fg, bg in segments:
            text = text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}').replace('\n', '\\par\n')
            
            style_cmd = ""
            if bg: style_cmd += f"\\highlight{color_idx[bg]}"
            if fg: style_cmd += f"\\cf{color_idx[fg]}\\b "
            
            if style_cmd:
                rtf.append(f"{{{style_cmd} {text}}}")
            else:
                rtf.append(text)

        rtf.append("}")
        with open(file_path, "w") as f: f.write("".join(rtf))
        messagebox.showinfo("Success", "Saved RTF.")

if __name__ == "__main__":
    root = tk.Tk()
    app = MufasaV2(root)
    root.mainloop()